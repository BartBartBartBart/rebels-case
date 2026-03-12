from pathlib import Path
from docx import Document
from sqlalchemy.orm import Session

from db_model import Doc
from services.classifier import BaseClassifier


def scan_folder(folder_path: str) -> list[Document]:
    """
    Recursively extract all .docx files in the specified folder path.

    Args:
        folder_path (str): The path to the folder to be scanned.
    Returns:
        List[Document]: A list of Document objects representing the .docx
            files found in the folder.
    """
    files = []

    # Recursively scan for .docx files in dir and subdirs
    for file in Path(folder_path).rglob("*.*"):

        # Docx files
        if file.is_file() and file.suffix == ".docx":
            print(f"Found document: {file}")
            files.append((Document(file), file))

        # Other files
        elif file.is_file():
            print(f"Skipping non-docx file: {file}")

    return files


def upsert_doc_metadata(db: Session, metadata: dict):
    """
    Upsert document metadata into the database. If a document with the same
    filename already exists, it will be updated with the new metadata. If not,
    a new entry will be created.

    Args:
        db (Session): The database session to use for the operation.
        metadata (dict): A dictionary containing the document metadata to be
            upserted.
    """
    existing_doc = db.query(Doc).filter(Doc.filename == metadata["filename"]).first()

    if existing_doc:
        # Update existing entry
        for key, value in metadata.items():
            setattr(existing_doc, key, value)
        print(f"Updated existing document: {metadata['filename']}")
    else:
        # Create new entry
        new_doc = Doc(**metadata)
        db.add(new_doc)
        print(f"Added new document: {metadata['filename']}")


def extract_metadata(doc: Document, file_path: str) -> dict:
    """
    Extract metadata from a Document object.

    Args:
        doc (Document): The Document object to extract metadata from.
        file_path (str): The file path of the document, used to extract
            filename and file type.
    Returns:
        dict: A dictionary containing the extracted metadata.
    """
    props = doc.core_properties
    metadata = {
        "filename": Path(file_path).name,
        "author": props.author,
        "created": props.created,
        "modified": props.modified,
        "file_type": Path(file_path).suffix,
        "file_size": Path(file_path).stat().st_size,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "word_count": sum(len(p.text.split()) for p in doc.paragraphs),
    }

    return metadata


def get_folder_insights(folder_path: str, db: Session) -> dict:
    """
    Ingest documents from the specified folder path and store
    their metadata in the database. This function reads all .docx
    files in the given folder, extracts their metadata, and saves
    it to the database.

    Args:
        folder_path (str): The path to the folder containing
            the documents to be ingested.
        db (Session): The database session to use for storing
            the metadata.
    Returns:
        dict: A dictionary containing the status of the ingestion
            process, the number of files ingested, and insights about
            the ingested documents.
    """

    print(f"Ingesting documents from folder: {folder_path}")

    # Find all files in folder
    files = scan_folder(folder_path)

    # Extract metadata and save to database
    insights = {}
    for file, file_path in files:
        metadata = extract_metadata(file, file_path=file_path)
        upsert_doc_metadata(db, metadata)
        insights[metadata["filename"]] = metadata

    db.commit()

    return {
        "num_files": len(files),
        "insights": insights,
    }


def get_folder_classifications(
    classifier: BaseClassifier, folder_path: str, overwrite: bool, db: Session
) -> dict:
    """
    Finds all .docx files in the folder and performs document classification
    using the given classifier in batches. Does not classify if a label
    already exists in the database.

    Args:
        classifier (BaseClassifier): A document classifier. Can either
            be a zero-shot classifier from HF or Phi-4-mini-instruct from HF.
        folder_path (str): The path for the folder.
        overwrite (bool): A boolean denoting whether to perform classification
            on perviously classified documents (for debugging mostly).
        db (Session): The session for the database.
    Returns:
        dict: A dictionary containing the number of classified files,
            a dictionary of all newly classified files and their labels, and
            a dictionary of all files in the directory that were already labelled.
    """

    print(f"Classifying documents in folder: {folder_path}")

    # Find all files in folder
    files = scan_folder(folder_path)
    filenames = []
    texts = []
    already_classified = {}

    for file in files:
        filename = Path(file[1]).name
        doc_entry = db.query(Doc).filter(Doc.filename == filename).first()

        if not overwrite and doc_entry:
            if doc_entry.label == "unlabeled":
                text = "\n".join(p.text for p in file[0].paragraphs)
                filenames.append(filename)
                texts.append(text)
                pass
            else:
                already_classified[filename] = doc_entry.label
                print(
                    f"Document {filename} already classified as {doc_entry.label}. "
                    f"Skipping classification."
                )
        elif not overwrite:
            # New document, classify and add without metadata
            text = "\n".join(p.text for p in file[0].paragraphs)
            filenames.append(filename)
            texts.append(text)
        else:
            text = "\n".join(p.text for p in file[0].paragraphs)
            filenames.append(filename)
            texts.append(text)

    # Single pass w/ batching
    if len(texts) > 0:
        pred_labels = classifier.classify_batch(texts)
    else:
        pred_labels = []

    for file, label in zip(filenames, pred_labels):
        upsert_doc_metadata(db, {"filename": file, "label": label})

    db.commit()

    return {
        "num_files": len(files),
        "new_classifications": {
            filename: label for filename, label in zip(filenames, pred_labels)
        },
        "already_classified": already_classified,
    }
