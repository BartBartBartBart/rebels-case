from pathlib import Path
from docx import Document
from pypdf import PdfReader
from datetime import datetime
import langdetect
from sqlalchemy.orm import Session

from db_model import Doc
from services.classifier import BaseClassifier
from logger import get_logger

logger = get_logger("services.insights")


def read(file: Path) -> str:
    """
    Extracts the text from .docx, .pdf and .txt files.

    Args:
        file (Path): Path leading to the file.
    Returns:
        str: Contents of the file.
    """

    if file.suffix.lower() == ".docx":
        return " ".join([p.text for p in Document(file).paragraphs])
    elif file.suffix.lower() == ".pdf":
        reader = PdfReader(file)
        return " ".join([page.extract_text() for page in reader.pages])
    elif file.suffix.lower() == ".txt":
        return file.read_text(encoding="utf-8")
    else:
        logger.waning(f"Can't read unsupported file {file.suffix.lower()}.")
        raise ValueError(f"Can't read unsupported file {file.suffix.lower()}.")


def scan_folder(folder_path: str) -> list[Path]:
    """
    Recursively extract all .docx files in the specified folder path.

    Args:
        folder_path (str): The path to the folder to be scanned.
    Returns:
        List[Document]: A list of Document objects representing the .docx
            files found in the folder.
    """
    supported_files = [".docx", ".pdf", ".txt"]
    files = []

    # Recursively scan for files in dir and subdirs
    for file in Path(folder_path).rglob("*.*"):

        if file.is_file() and file.suffix.lower() in supported_files:
            logger.info(f"Found document: {file}.")
            files.append(file)

        # Other files
        elif file.is_file():
            logger.info(f"Skipping unsupported file: {file}")

    return files


def upsert_doc_metadata(db: Session, metadata: dict):
    """
    Upsert document metadata into the database. If a document with the same
    filename already exists, it will be updated with the new metadata. If not,
    a new entry will be created.

    Args:
        db (Session): The database session to use for the operation.
        metadata (dict): A dictionary containing the document metadata to be
            upserted.
    """
    existing_doc = db.query(Doc).filter(Doc.filename == metadata["filename"]).first()

    if existing_doc:
        # Update existing entry
        for key, value in metadata.items():
            setattr(existing_doc, key, value)
        logger.info(f"Updated existing document: {metadata['filename']}")
    else:
        # Create new entry
        new_doc = Doc(**metadata)
        db.add(new_doc)
        logger.info(f"Added new document: {metadata['filename']}")


def extract_docx(file: Path) -> dict:
    """
    Extract metadata from .docx files.

    Args:
        file (Path): Path leading to .docx file.
    Returns:
        dict: Dictionary containing metadata.
    """

    doc = Document(file)
    props = doc.core_properties
    words = []
    for p in doc.paragraphs:
        words.extend(p.text.split(" "))

    return {
        "author": props.author,
        "created": props.created,
        "modified": props.modified,
        "language": langdetect.detect(" ".join(words)) if words else None,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "word_count": len(words),
    }


def extract_pdf(file: Path) -> dict:
    """
    Extract metadata from .pdf files.

    Args:
        file (Path): Path leading to .pdf file.
    Returns:
        dict: Dictionary containing metadata.
    """

    reader = PdfReader(file)
    words = []
    for page in reader.pages:
        words.extend(page.extract_text().split(" "))

    return {
        "author": reader.metadata.author,
        "title": reader.metadata.title,
        "created": reader.metadata.creation_date,
        "modified": reader.metadata.modification_date,
        "keywords": "".join(reader.metadata.keywords or []),
        "subject": reader.metadata.subject,
        "language": langdetect.detect(" ".join(words)) if words else None,
        "created_with": reader.metadata.creator,
        "word_count": len(words),
    }


def extract_txt(file: Path) -> dict:
    """
    Extract metadata from .txt files.

    Args:
        file (Path): Path leading to .txt file.
    Returns:
        dict: Dictionary containing metadata.
    """

    stat = file.stat()
    text = file.read_text(encoding="utf-8")
    words = text.split()

    return {
        "created": datetime.fromtimestamp(stat.st_ctime),
        "modified": datetime.fromtimestamp(stat.st_mtime),
        "word_count": len(words),
        "language": langdetect.detect(" ".join(words) if words else None),
    }


def extract_metadata(file: Path) -> dict:
    """
    Extract metadata from a Document object.

    Args:
        doc (Document): The Document object to extract metadata from.
        file_path (str): The file path of the document, used to extract
            filename and file type.
    Returns:
        dict: A dictionary containing the extracted metadata.
    """

    metadata = {
        "filename": Path(file).name,
        "file_type": Path(file).suffix,
        "file_size": Path(file).stat().st_size,
    }

    if file.suffix.lower() == ".docx":
        metadata.update(extract_docx(file))
    elif file.suffix.lower() == ".pdf":
        metadata.update(extract_pdf(file))
    elif file.suffix.lower() == ".txt":
        metadata.update(extract_txt(file))

    # Ensure we never store empty values for any file type.
    metadata = {k: v for k, v in metadata.items() if v is not None and v != ""}

    return metadata


def get_folder_insights(folder_path: str, db: Session) -> dict:
    """
    Ingest documents from the specified folder path and store
    their metadata in the database. This function reads all .docx
    files in the given folder, extracts their metadata, and saves
    it to the database.

    Args:
        folder_path (str): The path to the folder containing
            the documents to be ingested.
        db (Session): The database session to use for storing
            the metadata.
    Returns:
        dict: A dictionary containing the status of the ingestion
            process, the number of files ingested, and insights about
            the ingested documents.
    """

    logger.info(f"Ingesting documents from folder: {folder_path}")

    # Find all files in folder
    files = scan_folder(folder_path)

    # Extract metadata and save to database
    insights = {}
    for file in files:
        filename = Path(file).name
        doc_entry = db.query(Doc).filter(Doc.filename == filename).first()

        if doc_entry:
            # Extract metadata from database instead
            logger.info("Extracting metadata from database.")
            metadata = {}
            for col in Doc.__table__.columns:
                if getattr(doc_entry, col.name) is not None:
                    metadata[col.name] = getattr(doc_entry, col.name)
        else:
            # Extract metadata from file
            logger.info("Extracting metadata from file.")
            metadata = extract_metadata(file=file)
            upsert_doc_metadata(db, metadata)

        insights[metadata["filename"]] = metadata

    db.commit()

    return {
        "num_files": len(files),
        "insights": insights,
    }


def get_folder_classifications(
    classifier: BaseClassifier, folder_path: str, overwrite: bool, db: Session
) -> dict:
    """
    Finds all .docx files in the folder and performs document classification
    using the given classifier in batches. Does not classify if a label
    already exists in the database.

    Args:
        classifier (BaseClassifier): A document classifier. Can either
            be a zero-shot classifier from HF or Phi-4-mini-instruct from HF.
        folder_path (str): The path for the folder.
        overwrite (bool): A boolean denoting whether to perform classification
            on perviously classified documents (for debugging mostly).
        db (Session): The session for the database.
    Returns:
        dict: A dictionary containing the number of classified files,
            a dictionary of all newly classified files and their labels, and
            a dictionary of all files in the directory that were already labelled.
    """

    logger.info(f"Classifying documents in folder: {folder_path}")

    # Find all files in folder
    files = scan_folder(folder_path)
    filenames = []
    texts = []
    already_classified = {}

    for file in files:
        filename = Path(file).name
        doc_entry = db.query(Doc).filter(Doc.filename == filename).first()

        if not overwrite and doc_entry:
            if doc_entry.label == "unlabeled":
                filenames.append(filename)
                texts.append(read(file))
            else:
                already_classified[filename] = doc_entry.label
                logger.info(
                    f"Document {filename} already classified as {doc_entry.label}. "
                    f"Skipping classification."
                )
        elif not overwrite:
            # New document, classify and add without metadata
            filenames.append(filename)
            texts.append(read(file))
        else:
            filenames.append(filename)
            texts.append(read(file))

    # Single pass w/ batching
    if len(texts) > 0:
        pred_labels = classifier.classify_batch(texts)
    else:
        pred_labels = []

    for filename, label in zip(filenames, pred_labels):
        upsert_doc_metadata(db, {"filename": filename, "label": label})

    db.commit()

    classifications = {
        filename: label for filename, label in zip(filenames, pred_labels)
    }
    classifications.update(already_classified)

    return {"num_files": len(files), "classifications": classifications}
